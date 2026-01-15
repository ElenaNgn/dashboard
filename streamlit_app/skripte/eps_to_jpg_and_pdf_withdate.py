# -*- coding: utf-8 -*-
"""
Masszeichnungs-Verarbeitung mit Streamlit - MIT BACKGROUND PROCESSING
EPS/JPG → JPG + PDF + Excel (SAP)
Prozesse laufen im Hintergrund weiter, auch bei Seitenwechsel
"""

from pathlib import Path
from datetime import datetime
import shutil
import traceback
import json
import threading
import queue
from typing import Optional

import pandas as pd
from PIL import Image, ImageDraw, ImageFont, EpsImagePlugin
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx2pdf import convert
import streamlit as st


# ==========================================================
# KONFIGURATION
# ==========================================================

class Config:
    def __init__(self):
        self.eps_data = Path(r"D:\EPS")
        self.ziel_ordner = Path(r"D:\OUT")
        self.target_jpg_sap = Path(r"D:\SAP\YM1")
        self.target_pdf_sap = Path(r"D:\SAP\YM2")
        self.excel_output = Path(r"D:\SAP\import.xlsx")
        self.vorlage_docx = Path(r"D:\vorlage.docx")
        
        self.ghostscript = r"C:\Program Files\gs\gs10.03.0\bin\gswin64c.exe"
        
        self.dpi = 300
        self.width_cm = 17.3
        self.height_cm = 21.7
        self.jpg_quality = 97
        self.date_margin = 10
        
        self.font_path = "arial.ttf"
        self.font_size = 36


# ==========================================================
# HILFSFUNKTIONEN
# ==========================================================

def cm_to_px(cm: float, dpi: int) -> int:
    return int(cm / 2.54 * dpi)


def load_font(cfg: Config):
    try:
        return ImageFont.truetype(cfg.font_path, cfg.font_size)
    except Exception:
        return ImageFont.load_default()


def load_image(path: Path) -> Image.Image:
    if path.suffix.lower() == ".eps":
        with Image.open(path) as img:
            img.load(scale=12)
            return img.convert("RGB")
    else:
        with Image.open(path) as img:
            return img.convert("RGB")


# ==========================================================
# BILDVERARBEITUNG
# ==========================================================

def create_layout_image(img: Image.Image, name: str, cfg: Config, font) -> Path:
    w = cm_to_px(cfg.width_cm, cfg.dpi)
    h = cm_to_px(cfg.height_cm, cfg.dpi)

    scale = h / img.height
    new_size = (int(img.width * scale), int(img.height * scale))
    img = img.resize(new_size, Image.Resampling.LANCZOS)

    canvas = Image.new("RGB", (w, h), "white")
    canvas.paste(img, (0, 0))

    draw = ImageDraw.Draw(canvas)
    date = datetime.now().strftime("%d.%m.%Y")
    tw, th = draw.textbbox((0, 0), date, font=font)[2:]
    draw.text(
        (w - tw - cfg.date_margin, h - th - cfg.date_margin),
        date,
        fill="black",
        font=font,
    )

    out = cfg.ziel_ordner / f"{name}.jpg"
    canvas.save(out, "JPEG", dpi=(cfg.dpi, cfg.dpi), quality=cfg.jpg_quality)
    return out


# ==========================================================
# PDF
# ==========================================================

def create_pdf(image: Path, artikel: str, cfg: Config):
    tmp_docx = image.with_suffix(".tmp.docx")
    pdf_path = image.with_suffix(".pdf")

    try:
        doc = Document(cfg.vorlage_docx)

        p = doc.paragraphs[0]
        p.text = ""
        run = p.add_run(f"Artikelnummer {artikel[:4]} {artikel[4:]}")
        run.bold = True

        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(
            str(image),
            width=Mm(cfg.width_cm * 10)
        )

        doc.save(tmp_docx)
        convert(tmp_docx, pdf_path)

    except Exception as e:
        raise Exception(f"PDF-Erstellung fehlgeschlagen für {artikel}: {e}")

    finally:
        if tmp_docx.exists():
            try:
                tmp_docx.unlink()
            except Exception:
                pass


# ==========================================================
# BACKGROUND WORKER
# ==========================================================

class ProcessingJob:
    """Repräsentiert einen Verarbeitungsjob"""
    def __init__(self, job_id: str, cfg: Config):
        self.job_id = job_id
        self.cfg = cfg
        self.status = "queued"  # queued, processing, completed, error
        self.progress = 0.0
        self.current_file = ""
        self.total_files = 0
        self.processed_files = 0
        self.errors = []
        self.result_df: Optional[pd.DataFrame] = None
        self.start_time = datetime.now()
        self.end_time: Optional[datetime] = None


def process_job_worker(job: ProcessingJob, status_queue: queue.Queue):
    """Worker-Funktion die im Hintergrund läuft"""
    try:
        job.status = "processing"
        status_queue.put(("status", job))
        
        # Ghostscript konfigurieren
        EpsImagePlugin.gs_windows_binary = job.cfg.ghostscript
        Image.MAX_IMAGE_PIXELS = None
        
        # Zielordner erstellen
        job.cfg.ziel_ordner.mkdir(parents=True, exist_ok=True)
        
        font = load_font(job.cfg)
        files = {}
        
        # Dateien sammeln
        for p in job.cfg.eps_data.rglob("*"):
            if p.suffix.lower() in (".eps", ".jpg"):
                files.setdefault(p.stem, []).append(p)
        
        job.total_files = len(files)
        status_queue.put(("status", job))
        
        # Verarbeitung
        for idx, (name, variants) in enumerate(files.items()):
            try:
                job.current_file = name
                job.processed_files = idx
                job.progress = idx / job.total_files if job.total_files > 0 else 0
                status_queue.put(("status", job))
                
                src = next((p for p in variants if p.suffix.lower() == ".eps"), variants[0])
                img = load_image(src)
                jpg = create_layout_image(img, name, job.cfg, font)
                create_pdf(jpg, name, job.cfg)
                
            except Exception as e:
                error_msg = f"Fehler bei {name}: {str(e)}"
                job.errors.append(error_msg)
                status_queue.put(("error", error_msg))
        
        # Kopieren und Excel erstellen
        job.current_file = "Kopiere nach SAP..."
        status_queue.put(("status", job))
        
        job.cfg.target_jpg_sap.mkdir(parents=True, exist_ok=True)
        job.cfg.target_pdf_sap.mkdir(parents=True, exist_ok=True)
        
        jpgs = list(job.cfg.ziel_ordner.glob("*.jpg"))
        pdfs = list(job.cfg.ziel_ordner.glob("*.pdf"))
        
        for f in jpgs:
            shutil.copy2(f, job.cfg.target_jpg_sap / f.name)
        for f in pdfs:
            shutil.copy2(f, job.cfg.target_pdf_sap / f.name)
        
        # Excel erstellen
        job.current_file = "Erstelle Excel..."
        status_queue.put(("status", job))
        
        df = pd.DataFrame({
            "Reihenfolge": range(1, len(jpgs) + 1),
            "Artikel-Nr": [f.stem[:4] + " " + f.stem[4:] for f in jpgs],
            "Masszeichnung JPG (YM1)": [f"\\YM1\\{f.name}" for f in jpgs],
            "Masszeichnung PDF (YM2)": [f"\\YM2\\{f.stem}.pdf" for f in jpgs],
            "Status": ["allg. Mutation"] * len(jpgs)
        })
        
        df.to_excel(job.cfg.excel_output, index=False)
        job.result_df = df
        
        # Fertig!
        job.status = "completed"
        job.progress = 1.0
        job.end_time = datetime.now()
        job.current_file = "Abgeschlossen"
        status_queue.put(("status", job))
        status_queue.put(("complete", job))
        
    except Exception as e:
        job.status = "error"
        job.errors.append(f"Fataler Fehler: {str(e)}")
        status_queue.put(("status", job))
        status_queue.put(("fatal_error", str(e)))


# ==========================================================
# CONFIG LADEN
# ==========================================================

def load_config_from_json(path: Path) -> Config:
    data = json.loads(path.read_text(encoding="utf-8"))

    cfg = Config()
    cfg.eps_data = Path(data["eps_data"])
    cfg.ziel_ordner = Path(data["ziel_ordner"])
    cfg.vorlage_docx = Path(data["vorlage_docx"])

    cfg.target_jpg_sap = Path(
        data.get("target_jpg_sap", r"S:/Multimedia/SAP/YM1")
    )
    cfg.target_pdf_sap = Path(
        data.get("target_pdf_sap", r"S:/Multimedia/SAP/YM2")
    )

    cfg.excel_output = Path(
        data.get(
            "excel_output",
            Path(cfg.ziel_ordner) / "Import_MZ.xlsx"
        )
    )

    cfg.height_cm = 9.6 if data.get("halbseitig", False) else 21.7

    return cfg


# ==========================================================
# STREAMLIT UI
# ==========================================================

def main():
    st.title("📐 Masszeichnungs-Verarbeitung")
    st.write("EPS/JPG → JPG + PDF + Excel (SAP)")
    
    # Session State initialisieren
    if 'jobs' not in st.session_state:
        st.session_state.jobs = {}
    if 'status_queues' not in st.session_state:
        st.session_state.status_queues = {}
    
    st.sidebar.header("Konfiguration")
    
    # Konfiguration laden
    config_file = st.sidebar.file_uploader(
        "JSON-Konfiguration hochladen", 
        type=['json']
    )
    
    use_manual = st.sidebar.checkbox("Manuelle Konfiguration")
    
    cfg = None
    
    if config_file:
        try:
            temp_path = Path("temp_config.json")
            temp_path.write_bytes(config_file.read())
            cfg = load_config_from_json(temp_path)
            temp_path.unlink()
            st.sidebar.success("✅ Konfiguration geladen")
        except Exception as e:
            st.sidebar.error(f"Fehler beim Laden: {e}")
    
    elif use_manual:
        cfg = Config()
        st.sidebar.subheader("Pfade")
        cfg.eps_data = Path(st.sidebar.text_input(
            "EPS-Quellordner", 
            value=str(cfg.eps_data)
        ))
        cfg.ziel_ordner = Path(st.sidebar.text_input(
            "Zielordner", 
            value=str(cfg.ziel_ordner)
        ))
        cfg.vorlage_docx = Path(st.sidebar.text_input(
            "Word-Vorlage", 
            value=str(cfg.vorlage_docx)
        ))
        
        halbseitig = st.sidebar.checkbox("Halbseitiges Format")
        cfg.height_cm = 9.6 if halbseitig else 21.7
    
    # Neuen Job starten
    if cfg and st.button("🚀 Verarbeitung starten", type="primary"):
        job_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        job = ProcessingJob(job_id, cfg)
        status_queue = queue.Queue()
        
        # Job speichern
        st.session_state.jobs[job_id] = job
        st.session_state.status_queues[job_id] = status_queue
        
        # Thread starten (läuft im Hintergrund!)
        thread = threading.Thread(
            target=process_job_worker,
            args=(job, status_queue),
            daemon=True
        )
        thread.start()
        
        st.success(f"✅ Job {job_id} gestartet!")
        st.info("💡 Sie können jetzt zu anderen Seiten navigieren. Der Job läuft weiter!")
        st.rerun()
    
    # Status-Updates aus Queue holen
    for job_id, status_queue in st.session_state.status_queues.items():
        try:
            while True:
                msg_type, msg_data = status_queue.get_nowait()
                if msg_type in ["status", "complete"]:
                    st.session_state.jobs[job_id] = msg_data
        except queue.Empty:
            pass
    
    # Jobs anzeigen
    if st.session_state.jobs:
        st.header("📋 Aktive und abgeschlossene Jobs")
        
        for job_id, job in sorted(st.session_state.jobs.items(), reverse=True):
            with st.expander(
                f"Job {job_id} - {job.status.upper()}", 
                expanded=(job.status == "processing")
            ):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Status", job.status)
                with col2:
                    st.metric("Dateien", f"{job.processed_files}/{job.total_files}")
                with col3:
                    if job.end_time:
                        duration = (job.end_time - job.start_time).total_seconds()
                        st.metric("Dauer", f"{duration:.1f}s")
                    else:
                        st.metric("Läuft seit", f"{(datetime.now() - job.start_time).seconds}s")
                
                if job.status == "processing":
                    st.progress(job.progress)
                    st.text(f"Aktuell: {job.current_file}")
                
                if job.errors:
                    with st.expander(f"⚠️ {len(job.errors)} Fehler"):
                        for error in job.errors:
                            st.warning(error)
                
                if job.status == "completed" and job.result_df is not None:
                    st.success("✅ Erfolgreich abgeschlossen!")
                    st.subheader("📊 Excel-Vorschau")
                    st.dataframe(job.result_df)
                    st.info(f"Excel gespeichert: {job.cfg.excel_output}")
                    
                    if st.button(f"🗑️ Job {job_id} aus Liste entfernen", key=f"remove_{job_id}"):
                        del st.session_state.jobs[job_id]
                        if job_id in st.session_state.status_queues:
                            del st.session_state.status_queues[job_id]
                        st.rerun()
    
    elif not cfg:
        st.info("👈 Bitte Konfiguration laden oder manuell eingeben")
    
    # Auto-Refresh für laufende Jobs
    if any(job.status == "processing" for job in st.session_state.jobs.values()):
        import time
        time.sleep(2)
        st.rerun()


if __name__ == "__main__":
    main()